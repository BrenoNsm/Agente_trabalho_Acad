from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from datetime import datetime

class DocCreator:
    def __init__(self, nome_materia, nome_autor, sobre_o_documento, nome_curso, caminho_logo=None):
        self.nome_materia = nome_materia
        self.nome_autor = nome_autor
        self.sobre_o_documento = sobre_o_documento
        self.nome_curso = nome_curso
        self.caminho_logo = caminho_logo
        self.data_atual = datetime.now().strftime("%d/%m/%Y")
        self.doc = Document()
        self.__configurar_margens()
        self.__configurar_estilo()
        if self.caminho_logo:
            self.__adicionar_cabecalho_com_logo()

    def __configurar_margens(self):
        sections = self.doc.sections
        for section in sections:
            section.top_margin = Inches(1.18)    # 3 cm
            section.left_margin = Inches(1.18)   # 3 cm
            section.bottom_margin = Inches(0.79) # 2 cm
            section.right_margin = Inches(0.79)  # 2 cm

    def __configurar_estilo(self):
        style = self.doc.styles['Normal']
        style.font.name = 'Times New Roman'
        style.font.size = Pt(12)
        style.paragraph_format.line_spacing = 1.5

    def __adicionar_cabecalho_com_logo(self):
        section = self.doc.sections[0]
        header = section.header
        paragraph = header.paragraphs[0]
        run = paragraph.add_run()
        try:
            run.add_picture(self.caminho_logo, width=Inches(5.5))
        except Exception as e:
            print(f"[⚠️] Erro ao carregar imagem do cabeçalho: {e}")

    def _definir_idioma_ptbr(self, paragraph):
        if paragraph.runs:
            r = paragraph.runs[0]._element
            rPr = r.get_or_add_rPr()
            lang = OxmlElement('w:lang')
            lang.set(qn('w:val'), 'pt-BR')
            rPr.append(lang)

    def adicionar_capa(self):
        textos = [
            "UNIVERSIDADE ESTADUAL DE RORAIMA",  # Título fixo
            self.nome_materia,
            self.nome_autor,
            self.sobre_o_documento,
            self.nome_curso,
            self.data_atual
        ]

        for i, texto in enumerate(textos):
            par = self.doc.add_paragraph(texto)
            par.alignment = WD_ALIGN_PARAGRAPH.CENTER
            par.runs[0].bold = True
            self._definir_idioma_ptbr(par)

            # Garantir fonte, tamanho e espaçamento 1,15
            for run in par.runs:
                run.font.name = 'Times New Roman'
                run.font.size = Pt(12)

            par.paragraph_format.line_spacing = 1.15
            par.paragraph_format.space_after = Pt(0)  # Sem espaço extra entre parágrafos
            par.paragraph_format.space_before = Pt(0)


    def salvar(self, nome_arquivo="Trabalho_Academico_Gerado.docx"):
        self.doc.save(nome_arquivo)
