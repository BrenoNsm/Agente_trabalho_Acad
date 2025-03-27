import re
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

class MarkdownConverter:
    def __init__(self, doc: Document, permitir_listas=False):
        self.doc = doc
        self.permitir_listas = permitir_listas

    def adicionar_texto_formatado(self, markdown_texto: str):
        linhas = markdown_texto.split('\n')
        for linha in linhas:
            linha = linha.strip()
            if not linha:
                continue

            # Título H2 (## Título)
            if linha.startswith("## "):
                par = self.doc.add_paragraph()
                run = par.add_run(linha.replace("## ", "").strip())
                run.bold = True
                run.font.size = self.doc.styles['Normal'].font.size
                par.alignment = WD_ALIGN_PARAGRAPH.LEFT

            # Listas (* item) - só se permitido
            elif linha.startswith("* ") and self.permitir_listas:
                texto = linha.replace("*", "•").strip()
                par = self.doc.add_paragraph(f"{texto}")
                par.paragraph_format.left_indent = self.doc.styles['Normal'].paragraph_format.left_indent

            elif linha.startswith("* ") and not self.permitir_listas:
                # Transformar em parágrafo comum, sem bullet
                texto = linha.replace("*", "").strip()
                self.doc.add_paragraph(texto)

            # Negrito (**texto**)
            elif "**" in linha:
                par = self.doc.add_paragraph()
                partes = re.split(r"(\*\*.*?\*\*)", linha)
                for parte in partes:
                    run = par.add_run(parte.replace("**", ""))
                    if parte.startswith("**") and parte.endswith("**"):
                        run.bold = True
            else:
                self.doc.add_paragraph(linha)
